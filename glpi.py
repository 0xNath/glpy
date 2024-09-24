#!python
import requests
import base64
import html
import re
import sys
import argparse
import json

import atexit
import pytesseract
from io import BytesIO
from typing import BinaryIO

from openpyxl import load_workbook
from docx import Document as docxDocument
from pypdf import PdfReader
from PIL import Image
from bs4 import BeautifulSoup

class arguments:
    url = None
    username = ""
    password = ""
    userToken = ""
    sessionToken = ""
    applicationToken = ""
    APIPath = None

    searchInReplies = False
    searchInDOCX = False
    searchInXLSX = False
    searchInPDF = False
    searchInImages = False
    searchInUnknownFiles = False
    followTicketLinks = False

    ticketNumbers = None
    textToSearch = None
    closeSession = False


def setupArgumentsParsing():
    parser = argparse.ArgumentParser(
        prog="GLPI REST API client",
        description="This program can be used to retriview informations from a GLPI server.",
    )

    parser.add_argument("-url", "--url")
    parser.add_argument("-u", "--username", type=str)
    parser.add_argument("-p", "--password", type=str)
    parser.add_argument("-ut", "--userToken", type=str)
    parser.add_argument("-at", "--applicationToken", type=str)
    parser.add_argument("-ap", "--APIPath", default="/apirest.php", type=str)
    parser.add_argument("-st", "--sessionToken", type=str)
    parser.add_argument("-tn", "--ticketNumbers", type=str)
    parser.add_argument("-ts", "--textToSearch", type=str)
    parser.add_argument(
        "-sor",
        "--searchInReplies",
        default=False,
        action=argparse.BooleanOptionalAction,
    )
    parser.add_argument(
        "-sodx", "--searchInDOCX", default=False, action=argparse.BooleanOptionalAction
    )
    parser.add_argument(
        "-soxx", "--searchInXLSX", default=False, action=argparse.BooleanOptionalAction
    )
    parser.add_argument(
        "-sopf", "--searchInPDF", default=False, action=argparse.BooleanOptionalAction
    )
    parser.add_argument(
        "-soi", "--searchInImages", default=False, action=argparse.BooleanOptionalAction
    )
    parser.add_argument(
        "-suf",
        "--searchInUnknownFiles",
        default=False,
        action=argparse.BooleanOptionalAction,
    )
    parser.add_argument(
        "-ftl",
        "--followTicketLinks",
        default=False,
        action=argparse.BooleanOptionalAction,
    )
    parser.add_argument(
        "-cs", "--closeSession", default=False, action=argparse.BooleanOptionalAction
    )

    return parser.parse_args(namespace=arguments)


def verifyAuthentication(function):
    """Verify if the authentication against the server is done."""

    def verify(self, *args, **kargs):
        if not self._GLPI__authenticated:
            sys.stderr.write(
                str(
                    "You are not authenticated. "
                    "Use argument -h/--help , or look at the online documentation https://github.com/0xNath/glpy for help."
                )
            )
            sys.exit(1)

        return function(self, *args, **kargs)

    return verify


class GLPI:
    """A class made to facilitate scripting with python against the GLPI REST API."""

    def __init__(
        self,
        url: str,
        applicationToken: str,
        userToken: str = "",
        sessionToken: str = "",
        username: str = "",
        password: str = "",
        APIPath: str = "/apirest.php",
        closeSession: bool = True,
    ) -> None:
        atexit.register(self.__close)

        if len(url) == 0:
            raise (BaseException("You need to provide an URL."))

        self.url = url

        self.APIPath = APIPath

        self._closeSession = closeSession
        self.__authenticated = False

        self.session = requests.Session()
        self.session.headers = {"Content-Type": "application/json"}

        if len(applicationToken) > 0:
            self.setApplicationToken(applicationToken)

        if len(sessionToken) > 0:
            self._setSessionToken(args.sessionToken)
            self._closeSession = False
        elif len(userToken) > 0:
            self.authUsingToken(userToken)
        elif len(username) > 0 and len(password) > 0:
            self.authUsingCredentials(username, password)
        else:
            raise (
                BaseException(
                    "You need to provide either an user token or an username and a password to authenticate against the GLPI REST API."
                )
            )

    def __close(self) -> None:
        if self.__authenticated and self._closeSession:
            self.killSession()

    def _getSessionToken(self) -> str:
        return self.session.headers.get("Session-Token")

    def authUsingCredentials(self, username: str, password: str):
        credentials = username + ":" + password
        encodedCredentials = base64.b64encode(credentials.encode()).decode("utf-8")
        del credentials

        self.session.headers.update({"Authorization": f"Basic {encodedCredentials}"})
        del encodedCredentials

        self.__initSession()

    def authUsingToken(self, userToken: str):
        self.session.headers.update({"Authorization": f"user_token {userToken}"})
        self.__initSession()

    def setApplicationToken(self, appToken: str):
        self.session.headers.update({"App-Token": appToken})

    def _getApplicationToken(self) -> str:
        return self.session.headers.get("App-Token")

    def __initSession(self):
        initSessionRequest = self.session.get(
            f"{self.url}{self.APIPath}/initSession",
            params={"get_full_session": True},
        )

        if not initSessionRequest.status_code == 200:
            raise BaseException(
                str(initSessionRequest.status_code)
                + " "
                + str(initSessionRequest.reason)
                + ":\n"
                + str(initSessionRequest.json())
            )

        self.session.headers.pop("Authorization")

        self.session.headers.update(
            {
                "Session-Token": initSessionRequest.json()["session_token"],
            }
        )

        self.__authenticated = True

    def _setSessionToken(self, value):
        self.__authenticated = True
        self.session.headers.update({"Session-Token": value})

    def searchOptions(self, itemType: str = "AllAssets", raw: bool = False):

        parameters = {"raw": None} if raw else {}

        searchOptionsRequest = self.session.get(
            f"{self.url}{self.APIPath}/listSearchOptions/{itemType}",
            params=parameters,
        )

        if searchOptionsRequest.status_code == 200:
            result = {}

            for name, value in list(searchOptionsRequest.json().items()):
                if len(value) == 1:
                    last_name = name
                    result.update({name: {}})
                else:
                    buffer = dict(value)
                    buffer.update({"field": int(name)})
                    buffer_name = value["name"]
                    buffer.pop("name")

                    result[last_name].update({buffer_name: buffer})

            return result
        elif searchOptionsRequest.status_code == 401:
            raise searchOptionsRequest.json()
        else:
            raise BaseException(
                f"Unknow status code for this path : {searchOptionsRequest.status_code}.\n{searchOptionsRequest.json()}"
            )

    @verifyAuthentication
    def search(
        self,
        sort: int = 1,
        range: str = "0-49",
        rawdata: bool = False,
        withindexes: bool = False,
        uid_cols: bool = False,
        giveItems: bool = False,
        itemType: str = "AllAssets",
    ) -> dict:

        parameters = {
            "sort": sort,
            "range": range,
            "rawdata": rawdata,
            "withindexes": withindexes,
            "uid_cols": uid_cols,
            "giveItems": giveItems,
        }

        searchRequest = self.session.get(
            f"{self.url}{self.APIPath}/search/{itemType}", params=parameters
        )

        if searchRequest.status_code == 200 or searchRequest.status_code == 206:
            return searchRequest.json()
        elif searchRequest.status_code == 401:
            raise searchRequest.json()
        else:
            raise BaseException(
                f"Unknow status code for this path : {searchRequest.status_code}."
            )

    @verifyAuthentication
    def getItem(self, itemType: str, itemID: int) -> dict:

        getItemRequest = self.session.get(
            f"{self.url}{self.APIPath}/{itemType}/{itemID}"
        )

        if getItemRequest.status_code == 200 or getItemRequest.status_code == 206:
            return getItemRequest.json()
        else:
            sys.stderr.write(
                str(itemID)
                + " "
                + str(getItemRequest)
                + " "
                + getItemRequest.reason
                + "\n"
            )
            sys.exit(1)

    @verifyAuthentication
    def getSubItem(self, itemType: str, itemID: int, subItemType: str) -> dict:
        getItemRequest = self.session.get(
            f"{self.url}{self.APIPath}/{itemType}/{itemID}/{subItemType}"
        )

        return getItemRequest.json()

    @verifyAuthentication
    def getMyEntities(self, is_recursive: bool = False) -> str:

        parameters = {"is_recursive": is_recursive}

        getMyProfilesRequest = self.session.get(
            f"{self.url}{self.APIPath}/getMyProfiles", params=parameters
        )

        return getMyProfilesRequest.json()

    @verifyAuthentication
    def killSession(self) -> None:
        self.session.get(f"{self.url}{self.APIPath}/killSession")

    @verifyAuthentication
    def downloadDocument(self, documentID: int) -> BinaryIO:
        requestResponse = self.session.get(
            f"{self.url}{self.APIPath}/Document/{documentID}",
            headers={"Accept": "application/octet-stream"},
            stream=True,
        )

        if requestResponse.status_code == 200:
            return BytesIO(requestResponse.content)
        else:
            raise BaseException(requestResponse.json())

    def getTextFromImage(self, document: BinaryIO | bytes) -> str:
        if type(document) == bytes:
            document = BytesIO(document)
        return pytesseract.image_to_string(Image.open(document), lang="fre+eng")

    def parseAndSoupHTMLContent(self, text: str) -> str:
        return str(BeautifulSoup(html.unescape(text), features="html.parser"))

    def searchTextInItemDocuments(
        self,
        text: str,
        itemType: str,
        itemID: int,
        searchInImages: bool,
        searchInPDF: bool,
        searchInDOCX: bool,
        searchInXLSX: bool,
    ) -> tuple[bool, str]:
        for document in self.getSubItem(itemType, itemID, "Document"):
            if searchInImages and document["mime"].find("image") >= 0:
                textPosition = self.getTextFromImage(
                    self.downloadDocument(document["id"])
                ).find(text)

                if textPosition >= 0:
                    return (
                        True,
                        f"Found '{text}' in image '{document['name']}' at position {textPosition}.",
                    )
            elif searchInPDF and document["mime"] == "application/pdf":
                reader = PdfReader(self.downloadDocument(document["id"]))

                for pagePosition in range(0, len(reader.pages)):
                    page = reader.pages[pagePosition]
                    textPosition = page.extract_text().find(text)

                    if textPosition >= 0:
                        return (
                            True,
                            f"Found '{text}' in pdf '{document['name']}' at page {pagePosition}, position {textPosition}.",
                        )

                    for image in page.images:
                        textPosition = self.getTextFromImage(image.data).find(text)

                        if textPosition >= 0:
                            return (
                                True,
                                f"Found '{text}' in pdf '{document['name']}' at page {pagePosition}, in image '{image.name}' at position {textPosition}.",
                            )
            elif searchInDOCX and document["filename"].find(".docx") >= 0:
                reader = docxDocument(self.downloadDocument(document["id"]))

                for p in reader.paragraphs:
                    textPosition = p.text.find(text)

                    if textPosition >= 0:
                        return (
                            True,
                            f"Found '{text}' in docx '{document['name']}', position {textPosition}.",
                        )
            elif searchInXLSX and document["filename"].find(".xlsx") >= 0:
                wb = load_workbook(self.downloadDocument(document["id"]))

                for sheetName in wb.sheetnames:
                    sheet = wb[sheetName]
                    for row in range(1, sheet.max_row + 1):
                        for column in range(1, sheet.max_column + 1):
                            textPosition = str(sheet.cell(row, column).value).find(text)

                            if textPosition >= 0:
                                return (
                                    True,
                                    f"Found '{text}' in spreadsheet '{document['name']}', position {textPosition}.",
                                )

        return (False, "")

    @verifyAuthentication
    def deepSearchInTicket(
        self,
        text: str,
        ticketNumber: int,
        searchInReplies: bool = True,
        searchInDOCX: bool = True,
        searchInXLSX: bool = True,
        searchInPDF: bool = True,
        searchInImages: bool = False,
        searchInUnknownFiles: bool = True,
        followTicketLinks: bool = True,
    ) -> tuple[bool, str]:

        if len(text) < 0:
            raise (Exception("Please provide some text to search."))

        ticket = self.getItem("Ticket", itemID=ticketNumber)

        try:
            textPosition = self.parseAndSoupHTMLContent(ticket["name"]).find(text)
        except TypeError as e:
            print(e, ticket)

        if textPosition >= 0:
            return (
                True,
                f"Found '{text}' in ticket {ticket['id']} title at position {textPosition}.",
            )

        textPosition = self.parseAndSoupHTMLContent(ticket["content"]).find(text)

        if textPosition >= 0:
            return (
                True,
                f"Found '{text}' in ticket {ticket['id']} description at position {textPosition}.",
            )

        foundInItemDocuments, documentSearchResult = self.searchTextInItemDocuments(
            itemType="Ticket",
            text=text,
            itemID=ticket["id"],
            searchInDOCX=searchInDOCX,
            searchInXLSX=searchInXLSX,
            searchInPDF=searchInPDF,
            searchInImages=searchInImages,
        )

        if foundInItemDocuments:
            return (
                True,
                f"Result found in ticket {ticketNumber} -> {documentSearchResult}",
            )

        if searchInReplies:
            for subItemCategory in [
                "ITILFollowup",
                "TicketTask",
                "TicketValidation",
                "ITILSolution",
            ]:
                for subItem in self.getSubItem("Ticket", ticketNumber, subItemCategory):
                    if subItemCategory == "TicketValidation":
                        textPosition = self.parseAndSoupHTMLContent(
                            subItem["comment_submission"]
                        ).find(text)
                    else:
                        textPosition = self.parseAndSoupHTMLContent(
                            subItem["content"]
                        ).find(text)

                    if textPosition >= 0:
                        return (
                            True,
                            f"Found '{text}' at position {textPosition} in {subItemCategory} {subItem['id']}.",
                        )

                    foundInItemDocuments, documentSearchResult = (
                        self.searchTextInItemDocuments(
                            itemType="Ticket",
                            text=text,
                            itemID=ticket["id"],
                            searchInDOCX=searchInDOCX,
                            searchInXLSX=searchInXLSX,
                            searchInPDF=searchInPDF,
                            searchInImages=searchInImages,
                        )
                    )

                    if foundInItemDocuments:
                        return (
                            True,
                            f"Result found in ticket {ticketNumber} -> {documentSearchResult}",
                        )

        return (False, "")


if __name__ == "__main__":
    args = setupArgumentsParsing()

    server = GLPI(
        args.url,
        args.applicationToken,
        args.userToken,
        args.sessionToken,
        args.username,
        args.password,
    )

    for ticketNumber in args.ticketNumbers.split(","):
        try:
            isFound, message = server.deepSearchInTicket(
                text=args.textToSearch,
                ticketNumber=int(ticketNumber),
                searchInReplies=args.searchInReplies,
                searchInDOCX=args.searchInDOCX,
                searchInXLSX=args.searchInXLSX,
                searchInPDF=args.searchInPDF,
                searchInImages=args.searchInImages,
                searchInUnknownFiles=args.searchInUnknownFiles,
                followTicketLinks=args.followTicketLinks,
            )

            if isFound:
                sys.stdout.write(
                    json.dumps(
                        {
                            "Ticket.id": ticketNumber,
                            "progress": "Found",
                            "result": message,
                        }
                    )
                )
            else:
                sys.stdout.write(
                    json.dumps(
                        {
                            "Ticket.id": ticketNumber,
                            "progress": "Not found",
                            "result": "",
                        }
                    )
                )

            sys.stdout.flush()
        except Exception as e:
            sys.stderr.write(
                json.dumps(
                    {
                        "Ticket.id": ticketNumber,
                        "progress": "Error",
                        "result": str(e),
                    }
                )
            )
            sys.stderr.flush()

    sys.exit(0)
